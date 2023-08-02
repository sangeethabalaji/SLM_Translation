import os
import openai
import nltk
import re
import shutil
import time
import docx

from google.cloud import translate_v2 as translate
from google.transliteration import transliterate_text
from docx import Document
openai.api_key = "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"

# Set your Google Cloud project credentials path
# Make sure to replace "path/to/credentials.json" with the actual path to your credentials file
credentials_path = r"C:\Users\sange\Downloads\translate-389509-ef1d14dc7a3a.json"



#translate the given text in target language
def translate_text(text, target_language):
    translate_client = translate.Client.from_service_account_json(credentials_path)
    result = translate_client.translate(text, target_language=target_language)
    return result['translatedText']


#Read the given file line by line perform translation in the target language, and write it in output file
def translate_content(source_file_path, destination_file_path):
   
    try:
        #open file in read mode
        # Load the input .docx file
        doc = docx.Document(source_file_path)

        # Create a new Document object to store the contents
        new_doc = docx.Document()
        
        # Iterate through paragraphs in the input document and add them to the new document
        for para in doc.paragraphs:
            #Translate the given line in target language
            Trans_input = translate_text(para.text, target_language)
            
            #Replace &quot; with ""
            Final_Trans_input = replace_html_entities(Trans_input)
            
            new_doc.add_paragraph(Final_Trans_input)
        
        # Save the new document to the output .docx file
        new_doc.save(destination_file_path)
        
        print("File copied successfully!")
    except FileNotFoundError:
        print("One of the files was not found.")
    except IOError:
        print("An error occurred while copying the file.")


# Replace the HTML entity &quot; with a regular double quotation mark "
def replace_html_entities(text):
    return text.replace("&quot;", '"')

if __name__ == "__main__":

	# input and destination file paths
	input_file_path = r"C:\Users\sange\Downloads\test\word\SLM_Trial.docx"
	#destination folder in which output files to be written
	destination_file_path=r"C:\Users\sange\Downloads\test\output\Tamil_SLM_Trial.docx"

	#set the target language for translation
	target_language='ta'

	#call to the function for translation
	translate_content(input_file_path, destination_file_path)
